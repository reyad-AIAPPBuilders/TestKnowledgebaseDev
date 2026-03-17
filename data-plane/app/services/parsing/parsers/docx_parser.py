"""DOCX parser using python-docx."""

from docx import Document as DocxDocument

from app.services.parsing.models import DocumentMetadata, DocumentType, ParseOptions, TableData
from app.services.parsing.parsers.base import BaseParser, ParsedContent
from app.utils.logger import get_logger

log = get_logger(__name__)


class DocxParser(BaseParser):
    """Extract text from DOCX files."""

    def supports(self) -> list[str]:
        return [DocumentType.DOCX.value]

    async def parse(self, file_path: str, options: ParseOptions) -> ParsedContent:
        doc = DocxDocument(file_path)

        paragraphs: list[str] = []
        tables: list[TableData] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                if para.style and para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "").replace("Heading", "1")
                    try:
                        level_num = int(level)
                    except ValueError:
                        level_num = 1
                    paragraphs.append(f"{'#' * level_num} {text}")
                else:
                    paragraphs.append(text)

        if options.extract_tables:
            for table in doc.tables:
                try:
                    rows = []
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        rows.append(cells)

                    if len(rows) >= 2:
                        tables.append(TableData(headers=rows[0], rows=rows[1:]))
                    elif len(rows) == 1:
                        tables.append(TableData(headers=rows[0], rows=[]))
                except Exception as e:
                    log.debug("table_extraction_failed", error=str(e))

        full_text = "\n\n".join(paragraphs)

        core = doc.core_properties
        metadata = DocumentMetadata(
            title=core.title or None,
            author=core.author or None,
            subject=core.subject or None,
            creator=core.last_modified_by or None,
            creation_date=str(core.created) if core.created else None,
            modification_date=str(core.modified) if core.modified else None,
            word_count=len(full_text.split()) if full_text else 0,
            char_count=len(full_text),
        )

        return ParsedContent(
            text=full_text,
            metadata=metadata,
            tables=tables,
            pages_parsed=1,
            pages_failed=0,
        )
